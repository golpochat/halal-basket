#!/usr/bin/env python3
"""Generate Halal Basket stakeholder workflow DOCX + PDF + visual HTML."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "stakeholder-assets"
OUT_DOCX = ROOT / "Halal-Basket-Stakeholder-System-Workflow.docx"
OUT_PDF = ROOT / "Halal-Basket-Stakeholder-System-Workflow.pdf"
OUT_HTML = ROOT / "Halal-Basket-Stakeholder-System-Workflow.html"

GREEN = RGBColor(0x1B, 0x5E, 0x3B)
MUTED = RGBColor(0x4A, 0x55, 0x4A)

GLOSSARY = [
    ("Halal Basket", "The platform brand customers see. One catalogue experience, not separate shop storefronts."),
    ("Customer App", "Where shoppers browse, add items, check out, pay, and track orders."),
    ("Shop Portal", "Where a partner shop prepares orders, updates stock/prices, and assigns drivers."),
    ("Driver App", "Where drivers see assigned deliveries and mark progress through to delivered."),
    ("Admin Console", "Where operations and platform teams manage orders, users, catalogue, fees, legal, and WhatsApp."),
    ("Catalogue", "The shared product listing customers browse. Availability depends on the selected delivery area."),
    ("Basket / Cart", "The list of items the customer intends to buy, before the order is placed."),
    ("Delivery area", "A named place we deliver to (for example Lucan, Swords, Tallaght). Affects stock and delivery day."),
    ("Delivery calendar", "Rules that map each delivery area to the weekday(s) when deliveries run."),
    ("Scheduled delivery", "The live checkout option today: customer orders ahead for a planned delivery day."),
    ("Order", "The customer’s purchase record (items, totals, payment, delivery details)."),
    ("Fulfillment", "The shop’s part of an order — prepare, ready, out for delivery, delivered."),
    ("Stock hold", "A short reservation of stock while checkout completes, so items are not double-sold."),
    ("Payment", "Customer pays after placing the order (mock pay in pilot, or Stripe Checkout when configured)."),
    ("Coupon / Promo", "Discount codes or cart promotions applied at checkout (for example HALAL10, WELCOME5)."),
    ("Partner shop", "A store that stocks and prepares items for Halal Basket orders."),
    ("Warehouse", "A central stock location type in the system; not published for customer fulfillment in the current pilot."),
    ("WhatsApp channel", "Optional messaging path: order updates, customer care inbox, and catalogue cart → Halal Basket order."),
    ("Care inbox", "Admin view of WhatsApp conversations so support can reply and help customers."),
    ("Assist link", "A secure short-lived link that helps a customer continue checkout with support guidance."),
    ("Opt-in", "Customer agrees to receive WhatsApp order updates and share a phone number."),
    ("RBAC / Permissions", "Staff access controls — who can see ops, catalogue, GDPR, WhatsApp, and so on."),
    ("Platform vs Work", "Admin navigation: Platform = settings/governance; Work = day-to-day operations."),
    ("GDPR tools", "Admin privacy tools to search, export, or erase customer data under policy constraints."),
    ("Legal documents", "Published policies (privacy, terms, cookies, refunds) shown to customers."),
    ("Favourites", "Products a signed-in customer has hearted for quick re-find."),
    ("Live order status", "Customer order tracking updates while the order moves through shop and driver steps."),
]

ROLE_STEPS = [
    (
        "Customer",
        [
            "Register or sign in on the Customer App.",
            "Choose a delivery area and browse the catalogue.",
            "Add items to the basket (favourites optional).",
            "Checkout: confirm basket → delivery details → place order.",
            "Pay on the confirmation screen.",
            "Track the order under My orders; manage profile, addresses, and favourites.",
        ],
    ),
    (
        "Shop (partner)",
        [
            "Open the Shop Portal and review new fulfillments.",
            "Update status: preparing → ready → out for delivery.",
            "Assign a driver when ready for delivery.",
            "Maintain product price and stock quantity.",
            "Use the prep view to plan by delivery date.",
        ],
    ),
    (
        "Driver",
        [
            "Open the Driver App and review today’s assigned jobs.",
            "Open an order, complete the delivery steps, and mark progress.",
            "Submit feedback / ratings where prompted.",
            "Review history of completed work.",
        ],
    ),
    (
        "Admin / Operations",
        [
            "Look up orders, totals, coupons, and fees.",
            "Handle refunds, complaints, and customer risk/block actions when needed.",
            "Manage shops, drivers, staff users, and permissions.",
            "Maintain catalogue, featured categories, delivery fees, promotions, branding.",
            "Publish legal pages; use GDPR tools when required.",
            "Monitor WhatsApp care inbox and catalogue sync (when WhatsApp is enabled).",
        ],
    ),
    (
        "WhatsApp (optional channel)",
        [
            "Customer can opt in at checkout for order notifications.",
            "Keywords and inbox support help (status, pay link, order help).",
            "Support can send assist or shop links from the admin inbox.",
            "Customers can shop the WhatsApp catalogue; cart can create a Halal Basket order plus pay link.",
        ],
    ),
]

PILOT_LIVE = [
    "Customer browse → scheduled delivery order → pay → shop prepare → driver deliver → ops support",
    "Four portals: Customer, Shop, Driver, Admin",
    "Delivery areas & calendar (e.g. Lucan, Swords, Tallaght)",
    "Coupons / promotions, favourites, multi-language string packs (when languages are published)",
    "WhatsApp care + catalogue commerce in Meta sandbox / test setup",
]

PILOT_NOT_YET = [
    "Production / public staging website cutover (still local pilot-ready)",
    "Pickup and realtime delivery in the customer checkout screen (API exists; UI is scheduled-only today)",
    "Warehouse-led fulfillment published to customers",
    "Multi-shop split orders in production (feature exists, default off)",
    "WhatsApp on a live production business number (sandbox/test today)",
]


def set_run_font(run, *, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = GREEN
    return p


def add_para(doc: Document, text: str, *, bold=False, size=11, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=MUTED if not bold else GREEN)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11, color=MUTED)
        p.paragraph_format.space_after = Pt(4)


def add_image(doc: Document, path: Path, width_in: float = 6.3):
    if not path.exists():
        add_para(doc, f"[Diagram missing: {path.name}]", bold=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    p.paragraph_format.space_after = Pt(10)


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Halal Basket")
    set_run_font(r, size=28, bold=True, color=GREEN)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Stakeholder System Workflow")
    set_run_font(r, size=16, bold=True, color=MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "How the system works today · Pilot-ready overview · Glossary included\n"
        "Based on the live apps and features built so far"
    )
    set_run_font(r, size=10, color=MUTED)

    add_heading(doc, "1. What Halal Basket is", 1)
    add_para(
        doc,
        "Halal Basket is a grocery ordering and delivery platform. Customers shop one branded "
        "catalogue. Partner shops prepare the goods. Drivers deliver. Operations staff support "
        "orders, catalogue, and customer care — including WhatsApp when enabled.",
    )

    add_heading(doc, "2. Who uses the system", 1)
    add_image(doc, ASSETS / "hb-portals.png")
    add_para(
        doc,
        "Four apps work together. WhatsApp is a messaging channel that connects customers and "
        "support — it is not a separate login role.",
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(("Who", "App they use", "Main job")):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True, size=10, color=GREEN)
    rows = [
        ("Customer", "Customer App", "Browse, order, pay, track"),
        ("Shop staff", "Shop Portal", "Prepare orders, stock, assign driver"),
        ("Driver", "Driver App", "Deliver assigned orders"),
        ("Ops / Admin", "Admin Console", "Support, catalogue, settings, WhatsApp inbox"),
        ("Platform owner", "Admin Console (elevated)", "Roles, legal, branding, languages"),
    ]
    for who, app, job in rows:
        cells = table.add_row().cells
        cells[0].text = who
        cells[1].text = app
        cells[2].text = job
        for c in cells:
            for p in c.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10, color=MUTED)

    add_heading(doc, "3. How an order flows (happy path)", 1)
    add_image(doc, ASSETS / "hb-order-flow.png")
    add_para(doc, "The standard journey stakeholders should remember:", bold=True)
    add_bullets(
        doc,
        [
            "Customer browses the catalogue for their delivery area and builds a basket.",
            "Customer checks out for scheduled delivery and places the order (stock is held briefly).",
            "Customer pays on the confirmation screen.",
            "Shop prepares the fulfillment and assigns a driver.",
            "Driver delivers and updates status.",
            "Admin can look up the order, help with issues, refunds, or WhatsApp care if needed.",
        ],
    )

    add_heading(doc, "4. Who does what on one order", 1)
    add_image(doc, ASSETS / "hb-swimlanes.png")
    add_para(
        doc,
        "Think of one order as a relay: Customer starts it, Shop prepares it, Driver finishes "
        "delivery, Admin only steps in when support or configuration is needed.",
    )

    add_heading(doc, "5. Role workflows (step by step)", 1)
    for role, steps in ROLE_STEPS:
        add_heading(doc, role, 2)
        add_bullets(doc, steps)

    add_heading(doc, "6. Simple end-to-end picture", 1)
    add_para(
        doc,
        "Register → Catalogue (choose area) → Basket → Checkout (scheduled delivery) → "
        "Place order → Pay → Shop prepares & assigns driver → Driver delivers → "
        "Ops available for support",
        bold=True,
        size=11,
    )
    add_para(
        doc,
        "Optional side path: customer opts into WhatsApp for updates; or shops via WhatsApp "
        "catalogue, which can create the same unpaid order plus a pay link.",
    )

    add_heading(doc, "7. What is live for pilot vs not yet", 1)
    add_heading(doc, "Live / ready to demonstrate", 2)
    add_bullets(doc, PILOT_LIVE)
    add_heading(doc, "Not yet for stakeholders to treat as production-live", 2)
    add_bullets(doc, PILOT_NOT_YET)

    doc.add_page_break()
    add_heading(doc, "Part 2 — Glossary of terms", 1)
    add_para(
        doc,
        "Every important term used in this document is defined below. Use this section when "
        "someone asks “what does that mean?” during the meeting.",
    )
    gtable = doc.add_table(rows=1, cols=2)
    gtable.style = "Table Grid"
    ghdr = gtable.rows[0].cells
    ghdr[0].text = "Term"
    ghdr[1].text = "Meaning in Halal Basket"
    for cell in ghdr:
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True, size=10, color=GREEN)
    for term, meaning in GLOSSARY:
        cells = gtable.add_row().cells
        cells[0].text = term
        cells[1].text = meaning
        for p in cells[0].paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True, size=10, color=GREEN)
        for p in cells[1].paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, color=MUTED)

    add_heading(doc, "Meeting tip", 2)
    add_para(
        doc,
        "For a live walkthrough, open the visual companion HTML in a browser "
        "(Halal-Basket-Stakeholder-System-Workflow.html) and share your screen. "
        "Use this Word/PDF file as the leave-behind handout.",
    )

    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")


def build_pdf() -> None:
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="HBTitle",
            parent=styles["Title"],
            fontSize=22,
            textColor=colors.HexColor("#1B5E3B"),
            spaceAfter=6,
            alignment=TA_CENTER,
        )
    )
    styles.add(
        ParagraphStyle(
            name="HBSub",
            parent=styles["Normal"],
            fontSize=12,
            textColor=colors.HexColor("#4A554A"),
            alignment=TA_CENTER,
            spaceAfter=16,
        )
    )
    styles.add(
        ParagraphStyle(
            name="HBH1",
            parent=styles["Heading1"],
            fontSize=14,
            textColor=colors.HexColor("#1B5E3B"),
            spaceBefore=14,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="HBH2",
            parent=styles["Heading2"],
            fontSize=12,
            textColor=colors.HexColor("#1B5E3B"),
            spaceBefore=10,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="HBBody",
            parent=styles["Normal"],
            fontSize=10,
            textColor=colors.HexColor("#333933"),
            leading=14,
            alignment=TA_JUSTIFY,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="HBBullet",
            parent=styles["Normal"],
            fontSize=10,
            textColor=colors.HexColor("#333933"),
            leading=13,
            leftIndent=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="HBCaption",
            parent=styles["Normal"],
            fontSize=8,
            textColor=colors.HexColor("#6B736B"),
            alignment=TA_CENTER,
            spaceAfter=10,
        )
    )

    story = []
    story.append(Paragraph("Halal Basket", styles["HBTitle"]))
    story.append(Paragraph("Stakeholder System Workflow", styles["HBSub"]))
    story.append(
        Paragraph(
            "How the system works today · Pilot-ready overview · Glossary included<br/>"
            "Based on the live apps and features built so far",
            styles["HBSub"],
        )
    )

    story.append(Paragraph("1. What Halal Basket is", styles["HBH1"]))
    story.append(
        Paragraph(
            "Halal Basket is a grocery ordering and delivery platform. Customers shop one branded "
            "catalogue. Partner shops prepare the goods. Drivers deliver. Operations staff support "
            "orders, catalogue, and customer care — including WhatsApp when enabled.",
            styles["HBBody"],
        )
    )

    story.append(Paragraph("2. Who uses the system", styles["HBH1"]))
    img1 = ASSETS / "hb-portals.png"
    if img1.exists():
        story.append(RLImage(str(img1), width=6.2 * inch, height=3.5 * inch))
        story.append(Paragraph("Figure: Who uses Halal Basket", styles["HBCaption"]))

    role_data = [
        [
            Paragraph("<b>Who</b>", styles["HBBullet"]),
            Paragraph("<b>App</b>", styles["HBBullet"]),
            Paragraph("<b>Main job</b>", styles["HBBullet"]),
        ],
        ["Customer", "Customer App", "Browse, order, pay, track"],
        ["Shop staff", "Shop Portal", "Prepare orders, stock, assign driver"],
        ["Driver", "Driver App", "Deliver assigned orders"],
        ["Ops / Admin", "Admin Console", "Support, catalogue, settings, WhatsApp"],
        ["Platform owner", "Admin (elevated)", "Roles, legal, branding, languages"],
    ]
    t = Table(role_data, colWidths=[1.5 * inch, 1.7 * inch, 3.0 * inch])
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8F2EC")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#A8C4B4")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(t)
    story.append(Spacer(1, 8))

    story.append(Paragraph("3. How an order flows (happy path)", styles["HBH1"]))
    img2 = ASSETS / "hb-order-flow.png"
    if img2.exists():
        story.append(RLImage(str(img2), width=6.2 * inch, height=3.5 * inch))
        story.append(Paragraph("Figure: End-to-end order flow", styles["HBCaption"]))

    bullets = [
        "Customer browses the catalogue for their delivery area and builds a basket.",
        "Customer checks out for scheduled delivery and places the order (stock is held briefly).",
        "Customer pays on the confirmation screen.",
        "Shop prepares the fulfillment and assigns a driver.",
        "Driver delivers and updates status.",
        "Admin can look up the order, help with issues, refunds, or WhatsApp care if needed.",
    ]
    story.append(
        ListFlowable(
            [ListItem(Paragraph(b, styles["HBBullet"]), leftIndent=12, bulletColor=colors.HexColor("#1B5E3B")) for b in bullets],
            bulletType="bullet",
        )
    )

    story.append(Paragraph("4. Who does what on one order", styles["HBH1"]))
    img3 = ASSETS / "hb-swimlanes.png"
    if img3.exists():
        story.append(RLImage(str(img3), width=6.2 * inch, height=3.5 * inch))
        story.append(Paragraph("Figure: Role responsibilities across one order", styles["HBCaption"]))
    story.append(
        Paragraph(
            "Think of one order as a relay: Customer starts it, Shop prepares it, Driver finishes "
            "delivery, Admin only steps in when support or configuration is needed.",
            styles["HBBody"],
        )
    )

    story.append(Paragraph("5. Role workflows (step by step)", styles["HBH1"]))
    for role, steps in ROLE_STEPS:
        block = [Paragraph(role, styles["HBH2"])]
        block.append(
            ListFlowable(
                [ListItem(Paragraph(s, styles["HBBullet"]), leftIndent=12, bulletColor=colors.HexColor("#1B5E3B")) for s in steps],
                bulletType="bullet",
            )
        )
        story.append(KeepTogether(block))

    story.append(Paragraph("6. Simple end-to-end picture", styles["HBH1"]))
    story.append(
        Paragraph(
            "<b>Register → Catalogue (choose area) → Basket → Checkout (scheduled delivery) → "
            "Place order → Pay → Shop prepares &amp; assigns driver → Driver delivers → "
            "Ops available for support</b>",
            styles["HBBody"],
        )
    )
    story.append(
        Paragraph(
            "Optional side path: customer opts into WhatsApp for updates; or shops via WhatsApp "
            "catalogue, which can create the same unpaid order plus a pay link.",
            styles["HBBody"],
        )
    )

    story.append(Paragraph("7. What is live for pilot vs not yet", styles["HBH1"]))
    story.append(Paragraph("Live / ready to demonstrate", styles["HBH2"]))
    story.append(
        ListFlowable(
            [ListItem(Paragraph(x, styles["HBBullet"]), leftIndent=12, bulletColor=colors.HexColor("#1B5E3B")) for x in PILOT_LIVE],
            bulletType="bullet",
        )
    )
    story.append(Paragraph("Not yet for stakeholders to treat as production-live", styles["HBH2"]))
    story.append(
        ListFlowable(
            [ListItem(Paragraph(x, styles["HBBullet"]), leftIndent=12, bulletColor=colors.HexColor("#1B5E3B")) for x in PILOT_NOT_YET],
            bulletType="bullet",
        )
    )

    story.append(PageBreak())
    story.append(Paragraph("Part 2 — Glossary of terms", styles["HBH1"]))
    story.append(
        Paragraph(
            "Every important term used in this document is defined below.",
            styles["HBBody"],
        )
    )

    g_data = [
        [
            Paragraph("<b>Term</b>", styles["HBBullet"]),
            Paragraph("<b>Meaning in Halal Basket</b>", styles["HBBullet"]),
        ]
    ]
    for term, meaning in GLOSSARY:
        g_data.append(
            [
                Paragraph(f"<b>{term}</b>", styles["HBBullet"]),
                Paragraph(meaning, styles["HBBullet"]),
            ]
        )
    gt = Table(g_data, colWidths=[1.8 * inch, 4.4 * inch])
    gt.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8F2EC")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#A8C4B4")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7FAF8")]),
            ]
        )
    )
    story.append(gt)
    story.append(Spacer(1, 12))
    story.append(Paragraph("Meeting tip", styles["HBH2"]))
    story.append(
        Paragraph(
            "For a live walkthrough, open the visual companion HTML in a browser and share your "
            "screen. Use this PDF/Word file as the leave-behind handout.",
            styles["HBBody"],
        )
    )

    SimpleDocTemplate(
        str(OUT_PDF),
        pagesize=A4,
        leftMargin=0.7 * inch,
        rightMargin=0.7 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
        title="Halal Basket — Stakeholder System Workflow",
        author="Halal Basket",
    ).build(story)
    print(f"Wrote {OUT_PDF}")


def esc(s: str) -> str:
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def build_html() -> None:
    role_cards = ""
    for role, steps in ROLE_STEPS:
        lis = "".join(f"<li>{esc(s)}</li>" for s in steps)
        role_cards += f"""
        <section class="role-card">
          <h3>{esc(role)}</h3>
          <ol>{lis}</ol>
        </section>"""

    glossary_rows = "".join(
        f"<tr><th>{esc(t)}</th><td>{esc(m)}</td></tr>" for t, m in GLOSSARY
    )
    live_li = "".join(f"<li>{esc(x)}</li>" for x in PILOT_LIVE)
    not_li = "".join(f"<li>{esc(x)}</li>" for x in PILOT_NOT_YET)

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>Halal Basket — Stakeholder System Workflow</title>
  <style>
    :root {{
      --green: #1B5E3B;
      --green-soft: #E8F2EC;
      --green-mid: #3D7A5A;
      --ink: #243028;
      --muted: #5A665C;
      --line: #C5D6CB;
      --cream: #F4F8F5;
      --white: #ffffff;
      --warn-bg: #FFF8EF;
      --warn-border: #E6C48A;
      --ok-bg: #EEF7F1;
      --ok-border: #9CC9AE;
      --font: "Segoe UI", "Candara", "Calibri", system-ui, sans-serif;
      --display: "Georgia", "Palatino Linotype", "Book Antiqua", serif;
    }}
    * {{ box-sizing: border-box; }}
    html {{ scroll-behavior: smooth; }}
    body {{
      margin: 0;
      font-family: var(--font);
      color: var(--ink);
      background:
        radial-gradient(1200px 500px at 10% -10%, #d9eee2 0%, transparent 55%),
        radial-gradient(900px 400px at 100% 0%, #e7f1ea 0%, transparent 50%),
        linear-gradient(180deg, #f7fbf8 0%, #eef3f0 100%);
      line-height: 1.5;
    }}
    .wrap {{ max-width: 1080px; margin: 0 auto; padding: 28px 20px 80px; }}
    header.hero {{
      text-align: center;
      padding: 36px 16px 28px;
      margin-bottom: 28px;
      border-bottom: 1px solid var(--line);
    }}
    header.hero .brand {{
      font-family: var(--display);
      font-size: clamp(2rem, 4vw, 2.75rem);
      color: var(--green);
      margin: 0 0 6px;
      letter-spacing: -0.02em;
      animation: rise 0.7s ease both;
    }}
    header.hero h1 {{
      font-size: 1.15rem;
      font-weight: 600;
      color: var(--muted);
      margin: 0 0 10px;
      animation: rise 0.7s ease 0.08s both;
    }}
    header.hero p {{
      margin: 0 auto;
      max-width: 42rem;
      color: var(--muted);
      font-size: 0.98rem;
      animation: rise 0.7s ease 0.14s both;
    }}
    nav.toc {{
      display: flex;
      flex-wrap: wrap;
      gap: 8px;
      justify-content: center;
      margin: 22px 0 8px;
      animation: rise 0.7s ease 0.2s both;
    }}
    nav.toc a {{
      text-decoration: none;
      color: var(--green);
      background: var(--white);
      border: 1px solid var(--line);
      padding: 6px 12px;
      font-size: 0.85rem;
      border-radius: 6px;
      transition: background 0.2s, transform 0.2s;
    }}
    nav.toc a:hover {{ background: var(--green-soft); transform: translateY(-1px); }}
    section.block {{
      background: var(--white);
      border: 1px solid var(--line);
      border-radius: 14px;
      padding: 22px 22px 18px;
      margin: 18px 0;
      box-shadow: 0 8px 24px rgba(27, 94, 59, 0.05);
      animation: rise 0.6s ease both;
    }}
    section.block h2 {{
      font-family: var(--display);
      color: var(--green);
      margin: 0 0 12px;
      font-size: 1.35rem;
      font-weight: 600;
    }}
    section.block h3 {{
      color: var(--green-mid);
      margin: 16px 0 8px;
      font-size: 1.05rem;
    }}
    section.block p {{ color: var(--muted); margin: 0 0 12px; }}
    .figure {{
      margin: 8px 0 6px;
      border-radius: 10px;
      overflow: hidden;
      border: 1px solid var(--line);
      background: var(--cream);
    }}
    .figure img {{
      display: block;
      width: 100%;
      height: auto;
    }}
    .caption {{
      font-size: 0.8rem;
      color: var(--muted);
      text-align: center;
      margin: 6px 0 14px;
    }}
    /* Interactive flow strip */
    .flow {{
      display: grid;
      grid-template-columns: repeat(5, 1fr);
      gap: 10px;
      margin: 12px 0 8px;
    }}
    .flow-step {{
      background: linear-gradient(165deg, #f8fcf9, var(--green-soft));
      border: 1px solid var(--line);
      border-radius: 12px;
      padding: 14px 12px;
      text-align: center;
      position: relative;
      transition: transform 0.25s ease, box-shadow 0.25s ease;
      cursor: default;
    }}
    .flow-step:hover {{
      transform: translateY(-4px);
      box-shadow: 0 10px 20px rgba(27, 94, 59, 0.12);
    }}
    .flow-step .num {{
      width: 28px; height: 28px; line-height: 28px;
      margin: 0 auto 8px;
      border-radius: 50%;
      background: var(--green);
      color: white;
      font-weight: 700;
      font-size: 0.85rem;
    }}
    .flow-step strong {{
      display: block;
      color: var(--green);
      font-size: 0.92rem;
      margin-bottom: 4px;
    }}
    .flow-step span {{
      display: block;
      color: var(--muted);
      font-size: 0.78rem;
      line-height: 1.35;
    }}
    @media (max-width: 800px) {{
      .flow {{ grid-template-columns: 1fr 1fr; }}
    }}
    @media (max-width: 480px) {{
      .flow {{ grid-template-columns: 1fr; }}
    }}
    .roles {{
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(220px, 1fr));
      gap: 12px;
      margin-top: 8px;
    }}
    .role-pill {{
      background: var(--cream);
      border: 1px solid var(--line);
      border-radius: 12px;
      padding: 14px;
      transition: transform 0.2s;
    }}
    .role-pill:hover {{ transform: translateY(-2px); }}
    .role-pill h4 {{
      margin: 0 0 4px;
      color: var(--green);
      font-size: 1rem;
    }}
    .role-pill p {{ margin: 0; font-size: 0.85rem; }}
    .role-grid {{
      display: grid;
      grid-template-columns: 1fr;
      gap: 12px;
    }}
    @media (min-width: 720px) {{
      .role-grid {{ grid-template-columns: 1fr 1fr; }}
    }}
    .role-card {{
      border: 1px solid var(--line);
      border-radius: 12px;
      padding: 14px 16px;
      background: var(--cream);
    }}
    .role-card h3 {{ margin-top: 0; }}
    .role-card ol {{
      margin: 0;
      padding-left: 1.2rem;
      color: var(--muted);
      font-size: 0.9rem;
    }}
    .role-card li {{ margin: 4px 0; }}
    .chain {{
      font-family: var(--font);
      background: var(--green);
      color: #f4fff8;
      border-radius: 12px;
      padding: 16px 18px;
      font-size: 0.95rem;
      line-height: 1.55;
      letter-spacing: 0.01em;
    }}
    .split {{
      display: grid;
      grid-template-columns: 1fr;
      gap: 12px;
    }}
    @media (min-width: 720px) {{
      .split {{ grid-template-columns: 1fr 1fr; }}
    }}
    .panel {{
      border-radius: 12px;
      padding: 14px 16px;
      border: 1px solid;
    }}
    .panel.ok {{ background: var(--ok-bg); border-color: var(--ok-border); }}
    .panel.warn {{ background: var(--warn-bg); border-color: var(--warn-border); }}
    .panel h3 {{ margin: 0 0 8px; color: var(--ink); }}
    .panel ul {{ margin: 0; padding-left: 1.15rem; color: var(--muted); font-size: 0.9rem; }}
    .panel li {{ margin: 5px 0; }}
    table.glossary {{
      width: 100%;
      border-collapse: collapse;
      font-size: 0.9rem;
    }}
    table.glossary th, table.glossary td {{
      border-bottom: 1px solid var(--line);
      padding: 10px 8px;
      vertical-align: top;
      text-align: left;
    }}
    table.glossary th {{
      width: 28%;
      color: var(--green);
      font-weight: 700;
    }}
    table.glossary tr:nth-child(even) td,
    table.glossary tr:nth-child(even) th {{
      background: var(--cream);
    }}
    footer {{
      text-align: center;
      color: var(--muted);
      font-size: 0.85rem;
      margin-top: 28px;
    }}
    @keyframes rise {{
      from {{ opacity: 0; transform: translateY(12px); }}
      to {{ opacity: 1; transform: translateY(0); }}
    }}
    @media print {{
      body {{ background: white; }}
      nav.toc {{ display: none; }}
      section.block {{ box-shadow: none; break-inside: avoid; }}
      .flow-step:hover {{ transform: none; box-shadow: none; }}
    }}
  </style>
</head>
<body>
  <div class="wrap">
    <header class="hero">
      <p class="brand">Halal Basket</p>
      <h1>Stakeholder System Workflow</h1>
      <p>A clear picture of how the platform works today — who uses it, how an order moves, and what every term means.</p>
      <nav class="toc">
        <a href="#who">Who uses it</a>
        <a href="#flow">Order flow</a>
        <a href="#roles">Role steps</a>
        <a href="#status">Pilot status</a>
        <a href="#glossary">Glossary</a>
      </nav>
    </header>

    <section class="block" id="what">
      <h2>What Halal Basket is</h2>
      <p>Halal Basket is a grocery ordering and delivery platform. Customers shop one branded catalogue. Partner shops prepare the goods. Drivers deliver. Operations staff support orders, catalogue, and customer care — including WhatsApp when enabled.</p>
    </section>

    <section class="block" id="who">
      <h2>Who uses the system</h2>
      <div class="figure"><img src="stakeholder-assets/hb-portals.png" alt="Who uses Halal Basket — four portals plus WhatsApp channel" /></div>
      <p class="caption">Four apps work together. WhatsApp is a messaging channel, not a separate login role.</p>
      <div class="roles">
        <div class="role-pill"><h4>Customer</h4><p>Customer App — browse, order, pay, track</p></div>
        <div class="role-pill"><h4>Shop staff</h4><p>Shop Portal — prepare, stock, assign driver</p></div>
        <div class="role-pill"><h4>Driver</h4><p>Driver App — deliver assigned orders</p></div>
        <div class="role-pill"><h4>Admin / Ops</h4><p>Admin Console — support, settings, WhatsApp inbox</p></div>
      </div>
    </section>

    <section class="block" id="flow">
      <h2>How an order flows</h2>
      <div class="flow" aria-label="Order flow steps">
        <div class="flow-step"><div class="num">1</div><strong>Browse</strong><span>Catalogue by delivery area</span></div>
        <div class="flow-step"><div class="num">2</div><strong>Checkout</strong><span>Scheduled delivery + place order</span></div>
        <div class="flow-step"><div class="num">3</div><strong>Pay</strong><span>Confirmation screen payment</span></div>
        <div class="flow-step"><div class="num">4</div><strong>Prepare</strong><span>Shop fulfillment + assign driver</span></div>
        <div class="flow-step"><div class="num">5</div><strong>Deliver</strong><span>Driver completes; ops if needed</span></div>
      </div>
      <div class="figure" style="margin-top:16px"><img src="stakeholder-assets/hb-order-flow.png" alt="End-to-end order flow diagram" /></div>
      <p class="caption">Happy path from customer browse through to delivery and optional admin support.</p>
      <div class="figure"><img src="stakeholder-assets/hb-swimlanes.png" alt="Role swimlane responsibilities" /></div>
      <p class="caption">One order is a relay across Customer → Shop → Driver, with Admin on support.</p>
      <div class="chain">
        Register → Catalogue (choose area) → Basket → Checkout (scheduled delivery) → Place order → Pay → Shop prepares &amp; assigns driver → Driver delivers → Ops available for support
      </div>
      <p style="margin-top:12px">Optional side path: WhatsApp opt-in for updates, or WhatsApp catalogue cart creating the same order plus a pay link.</p>
    </section>

    <section class="block" id="roles">
      <h2>Role workflows (step by step)</h2>
      <div class="role-grid">{role_cards}</div>
    </section>

    <section class="block" id="status">
      <h2>What is live for pilot vs not yet</h2>
      <div class="split">
        <div class="panel ok">
          <h3>Ready to demonstrate</h3>
          <ul>{live_li}</ul>
        </div>
        <div class="panel warn">
          <h3>Not production-live yet</h3>
          <ul>{not_li}</ul>
        </div>
      </div>
    </section>

    <section class="block" id="glossary">
      <h2>Glossary of terms</h2>
      <p>Every important term used in this overview, in one place.</p>
      <table class="glossary">
        <thead><tr><th>Term</th><th>Meaning in Halal Basket</th></tr></thead>
        <tbody>{glossary_rows}</tbody>
      </table>
    </section>

    <footer>
      Halal Basket · Stakeholder pack · Open this file for screen-share; use the DOCX/PDF as the handout.
    </footer>
  </div>
</body>
</html>
"""
    OUT_HTML.write_text(html, encoding="utf-8")
    print(f"Wrote {OUT_HTML}")


if __name__ == "__main__":
    ASSETS.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_pdf()
    build_html()
    print("Done.")
